"""
report_module.py
Extracts telemetry dict from UnderwaterGaitAnalyzer instance and builds Word report (.docx).
"""
import os
import json
import subprocess
from datetime import datetime, timezone
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _safe_rom(vals):
    arr = np.asarray(vals, dtype=float)
    if arr.size == 0 or np.all(np.isnan(arr)):
        return 0.0
    return float(np.nanmax(arr) - np.nanmin(arr))


def _calc_correlation(l_vals, r_vals):
    l = np.asarray(l_vals, dtype=float)
    r = np.asarray(r_vals, dtype=float)
    mask = ~np.isnan(l) & ~np.isnan(r)
    if np.sum(mask) < 5:
        return 0.0
    corr = np.corrcoef(l[mask], r[mask])[0, 1]
    return 0.0 if np.isnan(corr) else round(float(corr), 2)


def _asymmetry_status(deg: float) -> str:
    if deg >= 20:
        return "High — clinical review recommended"
    if deg >= 10:
        return "Mild — monitor during rehab"
    return "Within normal range"


def _set_cell_shading(cell, hex_color: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shading)


def _style_table_header(table, headers: list[str]):
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        _set_cell_shading(hdr[i], "0B6E4F")


def _add_data_row(table, values: list[str], bold_first: bool = False):
    row = table.add_row().cells
    for i, val in enumerate(values):
        row[i].text = str(val)
        for p in row[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT
            if bold_first and i == 0:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)


def _build_clinical_summary(report_data: dict) -> list[str]:
    findings = []
    asym = report_data.get("asymmetry_deg", {})
    tracking = report_data.get("tracking", {})
    cadence = report_data.get("cadence_steps_per_min", 0)

    high_joints = [j for j, d in asym.items() if d >= 20]
    mild_joints = [j for j, d in asym.items() if 10 <= d < 20]

    if high_joints:
        joints_str = ", ".join(j.title() for j in high_joints)
        findings.append(f"Significant left-right asymmetry detected in {joints_str} (≥20°). Targeted intervention may be warranted.")

    if mild_joints:
        joints_str = ", ".join(j.title() for j in mild_joints)
        findings.append(f"Mild asymmetry noted in {joints_str} (10–20°). Continue monitoring across sessions.")

    if not high_joints and not mild_joints:
        findings.append("Bilateral joint ROM asymmetry is within acceptable clinical thresholds (<10°) for all measured joints.")

    if cadence < 80:
        findings.append(f"Cadence ({cadence:.0f} spm) is below typical walking range — may indicate reduced mobility or cautious gait.")
    elif cadence > 130:
        findings.append(f"Cadence ({cadence:.0f} spm) is elevated — assess for compensatory fast-stepping pattern.")
    else:
        findings.append(f"Cadence ({cadence:.0f} spm) falls within a typical functional walking range.")

    mean_conf = tracking.get("mean")
    if mean_conf is not None:
        pct = mean_conf * 100
        if pct < 50:
            findings.append(f"Pose tracking confidence is low ({pct:.0f}%) — interpret joint angles with caution; consider re-recording with better visibility.")
        elif pct < 70:
            findings.append(f"Pose tracking confidence is moderate ({pct:.0f}%) — results are usable but some frames may be unreliable.")

    return findings


def build_report_data(analyzer, session_label: str = "Session Report", **meta) -> dict:
    """Extract complete telemetry out of a finished UnderwaterGaitAnalyzer instance into a dictionary."""
    fps = analyzer.fps
    total_frames = analyzer.frame_count
    total_time = total_frames / fps if fps > 0 else 0.0
    total_steps = len(analyzer.step_frames_left) + len(analyzer.step_frames_right)
    cadence = (total_steps / total_time) * 60.0 if total_time > 0 else 0.0

    joints = ["knee", "hip", "ankle"]
    rom = {}
    mean_angles = {}
    correlation = {}

    for j in joints:
        l = np.asarray(analyzer.angles[f"left_{j}"], dtype=float)
        r = np.asarray(analyzer.angles[f"right_{j}"], dtype=float)
        rom[j] = {"left": round(_safe_rom(l), 2), "right": round(_safe_rom(r), 2)}
        mean_angles[j] = {
            "left": round(float(np.nanmean(l)), 2) if l.size and not np.all(np.isnan(l)) else None,
            "right": round(float(np.nanmean(r)), 2) if r.size and not np.all(np.isnan(r)) else None,
        }
        correlation[j] = _calc_correlation(l, r)

    asymmetry = {
        j: round(abs(rom[j]["left"] - rom[j]["right"]), 2) for j in joints
    }

    left_step_interval = None
    right_step_interval = None
    if len(analyzer.step_frames_left) > 1:
        intervals = np.diff(analyzer.step_frames_left) / fps
        left_step_interval = {"mean": round(float(np.mean(intervals)), 3),
                              "std": round(float(np.std(intervals)), 3)}
    if len(analyzer.step_frames_right) > 1:
        intervals = np.diff(analyzer.step_frames_right) / fps
        right_step_interval = {"mean": round(float(np.mean(intervals)), 3),
                               "std": round(float(np.std(intervals)), 3)}

    conf = np.asarray(analyzer.landmark_confidence, dtype=float)
    tracking = {
        "mean": round(float(np.mean(conf)), 4) if conf.size else None,
        "min": round(float(np.min(conf)), 4) if conf.size else None,
        "good_frames": int(np.sum(conf > 0.7)) if conf.size else 0,
        "total_frames": int(conf.size),
    }

    pelvis_tilt_rom = _safe_rom(analyzer.angles["pelvis_tilt"])
    pelvis_rot_rom = _safe_rom(analyzer.angles["pelvis_rotation"])
    pelvis_tilt_arr = np.asarray(analyzer.angles["pelvis_tilt"], dtype=float)
    mean_pelvis_tilt = round(float(np.nanmean(pelvis_tilt_arr)), 2) if pelvis_tilt_arr.size and not np.all(np.isnan(pelvis_tilt_arr)) else 0.0

    return {
        "session_label": session_label,
        "patient_name": meta.get("patient_name"),
        "recorded_date": meta.get("recorded_date"),
        "recorded_time": meta.get("recorded_time"),
        "generated_at": datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"),
        "total_frames": total_frames,
        "total_time_sec": round(total_time, 2),
        "fps": round(fps, 2),
        "total_steps": total_steps,
        "left_steps": len(analyzer.step_frames_left),
        "right_steps": len(analyzer.step_frames_right),
        "cadence_steps_per_min": round(cadence, 2),
        "rom": rom,
        "mean_angles": mean_angles,
        "asymmetry_deg": asymmetry,
        "correlation": correlation,
        "step_timing": {
            "left": left_step_interval,
            "right": right_step_interval,
        },
        "pelvis": {
            "tilt_rom_deg": round(pelvis_tilt_rom, 2),
            "rotation_rom_deg": round(pelvis_rot_rom, 2),
            "mean_tilt_deg": mean_pelvis_tilt,
        },
        "tracking": tracking,
    }


def generate_docx_report(
    report_data: dict,
    out_path: str,
    comprehensive_png: str = None,
    knee_png: str = None,
    hip_png: str = None,
    ankle_png: str = None,
):
    """Generate a clinician-friendly Word report with structured sections and embedded plots."""
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    js_script = Path(__file__).parent / "generate_docx.js"
    if js_script.exists():
        try:
            args = [
                "node", str(js_script),
                json.dumps(report_data),
                str(out_path),
                str(comprehensive_png or ""),
                str(knee_png or ""),
                str(hip_png or ""),
                str(ankle_png or "")
            ]
            res = subprocess.run(args, capture_output=True, text=True)
            if res.returncode == 0 and out_path.exists():
                return str(out_path)
        except Exception as e:
            print(f"⚠ Node script execution skipped: {e}")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Title block ──
    title = doc.add_heading("Underwater Gait Analysis", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(11, 110, 79)

    subtitle = doc.add_paragraph("Clinical Biomechanics Report")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in subtitle.runs:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    patient = report_data.get("patient_name") or "Patient"
    rec_date = report_data.get("recorded_date") or "—"
    rec_time = report_data.get("recorded_time") or ""
    session_label = report_data.get("session_label", "Session Report")

    info_table = doc.add_table(rows=4, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.style = "Table Grid"
    info_rows = [
        ("Patient", patient),
        ("Session", session_label),
        ("Recording Date / Time", f"{rec_date}  {rec_time}".strip()),
        ("Report Generated", report_data.get("generated_at", datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC"))),
    ]
    for i, (label, val) in enumerate(info_rows):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = val
        for p in info_table.rows[i].cells[0].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

    doc.add_paragraph()

    # ── Clinical summary ──
    doc.add_heading("Clinical Summary", level=1)
    summary_intro = doc.add_paragraph(
        "Automated analysis of underwater gait video using pose estimation. "
        "Findings below are intended to support clinical decision-making and should be interpreted alongside physical examination."
    )
    summary_intro.runs[0].font.italic = True
    summary_intro.runs[0].font.size = Pt(9)
    summary_intro.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    for finding in _build_clinical_summary(report_data):
        p = doc.add_paragraph(finding, style="List Bullet")
        p.runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # ── Gait parameters ──
    doc.add_heading("Gait Parameters", level=1)
    gait_table = doc.add_table(rows=1, cols=2)
    gait_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    gait_table.style = "Table Grid"
    _style_table_header(gait_table, ["Parameter", "Value"])

    tracking = report_data.get("tracking", {})
    mean_conf = tracking.get("mean")
    conf_str = f"{mean_conf:.1%}" if mean_conf else "N/A"

    gait_rows = [
        ("Recording Duration", f"{report_data['total_time_sec']} seconds"),
        ("Frame Rate", f"{report_data['fps']} FPS"),
        ("Total Frames Analyzed", str(report_data["total_frames"])),
        ("Total Steps Detected", str(report_data["total_steps"])),
        ("Left / Right Steps", f"{report_data['left_steps']} / {report_data['right_steps']}"),
        ("Cadence", f"{report_data['cadence_steps_per_min']} steps/min"),
        ("Tracking Confidence (mean)", conf_str),
        ("High-Quality Frames (>70%)", f"{tracking.get('good_frames', 0)} / {tracking.get('total_frames', 0)}"),
    ]
    for label, val in gait_rows:
        _add_data_row(gait_table, [label, val], bold_first=True)

    doc.add_paragraph()

    # ── ROM & asymmetry ──
    doc.add_heading("Range of Motion & Bilateral Comparison", level=1)
    doc.add_paragraph(
        "ROM = peak-to-peak joint angle during gait. Asymmetry thresholds: Normal <10°, Mild 10–20°, High ≥20°."
    ).runs[0].font.size = Pt(9)

    rom_table = doc.add_table(rows=1, cols=6)
    rom_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    rom_table.style = "Table Grid"
    _style_table_header(rom_table, [
        "Joint", "Left ROM (°)", "Right ROM (°)", "Asymmetry (°)", "L-R Correlation", "Clinical Status"
    ])

    for joint in ["knee", "hip", "ankle"]:
        asym = report_data["asymmetry_deg"][joint]
        status = _asymmetry_status(asym)
        _add_data_row(rom_table, [
            joint.title(),
            str(report_data["rom"][joint]["left"]),
            str(report_data["rom"][joint]["right"]),
            str(asym),
            str(report_data.get("correlation", {}).get(joint, "N/A")),
            status,
        ], bold_first=True)

    doc.add_paragraph()

    # ── Step timing ──
    step_timing = report_data.get("step_timing", {})
    if step_timing.get("left") or step_timing.get("right"):
        doc.add_heading("Step Timing", level=1)
        timing_table = doc.add_table(rows=1, cols=3)
        timing_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        timing_table.style = "Table Grid"
        _style_table_header(timing_table, ["Side", "Mean Interval (s)", "Std Dev (s)"])
        for side in ["left", "right"]:
            data = step_timing.get(side)
            if data:
                _add_data_row(timing_table, [side.title(), str(data["mean"]), str(data["std"])], bold_first=True)
        doc.add_paragraph()

    # ── Pelvis ──
    doc.add_heading("Pelvis Kinematics", level=1)
    p = report_data.get("pelvis", {})
    pelvis_table = doc.add_table(rows=1, cols=2)
    pelvis_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pelvis_table.style = "Table Grid"
    _style_table_header(pelvis_table, ["Metric", "Value"])
    for label, val in [
        ("Pelvic Tilt ROM", f"{p.get('tilt_rom_deg', 0)}°"),
        ("Pelvic Rotation ROM", f"{p.get('rotation_rom_deg', 0)}°"),
        ("Mean Pelvic Tilt", f"{p.get('mean_tilt_deg', 0)}°"),
    ]:
        _add_data_row(pelvis_table, [label, val], bold_first=True)

    doc.add_paragraph()

    # ── Visualizations ──
    doc.add_heading("Biomechanics Visualizations", level=1)
    doc.add_paragraph(
        "Figures generated from frame-by-frame joint angle analysis. Refer to annotated video for temporal context."
    ).runs[0].font.size = Pt(9)

    for name, img_path in [
        ("Comprehensive Gait Overview", comprehensive_png),
        ("Knee Joint Analysis", knee_png),
        ("Hip Joint Analysis", hip_png),
        ("Ankle Joint Analysis", ankle_png),
    ]:
        if img_path and os.path.exists(img_path):
            doc.add_heading(name, level=2)
            doc.add_picture(img_path, width=Inches(6.0))
            doc.add_paragraph()

    # ── Footer / disclaimer ──
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        "DISCLAIMER: This report was generated automatically by GaitRehab underwater gait analysis software. "
        "Results depend on video quality, patient positioning, and pose estimation accuracy. "
        "This document does not constitute a medical diagnosis and must be reviewed by a qualified clinician "
        "before informing treatment decisions."
    )
    disclaimer.runs[0].font.size = Pt(8)
    disclaimer.runs[0].font.italic = True
    disclaimer.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    doc.save(str(out_path))
    return str(out_path)
