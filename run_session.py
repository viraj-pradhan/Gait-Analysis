"""
run_session.py
Automated Session-Based Gait Report Generation CLI.

Usage:
  python run_session.py --video path/to/clip.mp4 [--patient "Name"] [--date YYYY-MM-DD] [--time HH:MM]
  python run_session.py --batch path/to/folder_of_clips/
  python run_session.py --trend [--since YYYY-MM-DD]
"""
import os
import sys
import shutil
import json
import argparse
import traceback
from pathlib import Path
from datetime import datetime, timezone

import numpy as np
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Import core modules
from underwater_gait_analyzer_fixed import UnderwaterGaitAnalyzer
from report_module import build_report_data, generate_docx_report

SESSIONS_DIR = Path(__file__).parent / "sessions"
INDEX_FILE = SESSIONS_DIR / "index.json"


# ── Registry Helper ────────────────────────────────────────────────────────────

def load_index() -> list:
    """Load registry array from sessions/index.json."""
    if not INDEX_FILE.exists():
        return []
    try:
        with open(INDEX_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"⚠ Warning: Could not parse index.json: {e}")
        return []


def save_index(index_data: list):
    """Save registry array to sessions/index.json."""
    SESSIONS_DIR.mkdir(parents=True, exist_ok=True)
    with open(INDEX_FILE, "w", encoding="utf-8") as f:
        json.dump(index_data, f, indent=2)


def update_index(meta_entry: dict):
    """Append or update a session entry in index.json."""
    index = load_index()
    session_id = meta_entry.get("session_id")
    updated = False
    for i, entry in enumerate(index):
        if entry.get("session_id") == session_id:
            index[i] = meta_entry
            updated = True
            break
    if not updated:
        index.append(meta_entry)

    save_index(index)


# ── Single Session Processing ─────────────────────────────────────────────────

def process_single_video(
    video_path: Path, 
    date_str: str = None, 
    patient_name: str = "Unknown Patient",
    time_str: str = None,
    progress_callback=None
) -> dict:
    """Process a single video into sessions/<date>/session_<N>/ directory."""
    video_path = Path(video_path).resolve()
    if not video_path.exists():
        raise FileNotFoundError(f"Input video file not found: {video_path}")

    now_dt = datetime.now()
    if not date_str:
        date_str = now_dt.strftime("%Y-%m-%d")
    if not time_str:
        time_str = now_dt.strftime("%H:%M")

    date_dir = SESSIONS_DIR / date_str
    date_dir.mkdir(parents=True, exist_ok=True)

    # Auto-increment session_N
    existing_sessions = [
        d for d in date_dir.iterdir()
        if d.is_dir() and d.name.startswith("session_")
    ]
    session_nums = []
    for d in existing_sessions:
        try:
            session_nums.append(int(d.name.split("_")[1]))
        except ValueError:
            pass

    next_num = max(session_nums) + 1 if session_nums else 1
    session_id = f"{date_str}/session_{next_num}"
    session_dir = date_dir / f"session_{next_num}"
    session_dir.mkdir(parents=True, exist_ok=True)

    session_label = f"Session {next_num} — {patient_name} ({date_str})"
    print(f"\n▶ Processing Session: {session_id} [{patient_name}] ({video_path.name})")

    # Copy input video
    input_video_dest = session_dir / "input.mp4"
    shutil.copy2(video_path, input_video_dest)

    annotated_output = session_dir / "output_annotated.mp4"

    meta_entry = {
        "session_id": session_id,
        "patient_name": patient_name,
        "recorded_date": date_str,
        "recorded_time": time_str,
        "video_filename": video_path.name,
        "date": date_str,
        "session_number": next_num,
        "session_label": session_label,
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "fps": 0.0,
        "duration_sec": 0.0,
        "cadence_spm": 0.0,
        "mean_confidence": 0.0,
        "status": "failed",
        "error": None,
        "session_path": str(session_dir.relative_to(Path(__file__).parent)),
        "report_docx": str((session_dir / "report.docx").relative_to(Path(__file__).parent)),
    }

    try:
        # Instantiate and run analyzer
        analyzer = UnderwaterGaitAnalyzer(
            video_path=str(input_video_dest),
            output_path=str(annotated_output),
            progress_callback=progress_callback
        )
        analyzer.run_complete_analysis(output_dir=str(session_dir))

        # Build telemetry report dictionary
        report_data = build_report_data(
            analyzer,
            session_label=session_label,
            patient_name=patient_name,
            recorded_date=date_str,
            recorded_time=time_str,
        )

        # File paths inside session directory
        comp_png = session_dir / "gait_analysis_comprehensive.png"
        knee_png = session_dir / "knee_analysis_detailed.png"
        hip_png = session_dir / "hip_analysis_detailed.png"
        ankle_png = session_dir / "ankle_analysis_detailed.png"
        report_docx = session_dir / "report.docx"
        report_json = session_dir / "report_data.json"

        # Generate Word report
        generate_docx_report(
            report_data=report_data,
            out_path=str(report_docx),
            comprehensive_png=str(comp_png) if comp_png.exists() else None,
            knee_png=str(knee_png) if knee_png.exists() else None,
            hip_png=str(hip_png) if hip_png.exists() else None,
            ankle_png=str(ankle_png) if ankle_png.exists() else None,
        )

        # Save JSON telemetry
        with open(report_json, "w", encoding="utf-8") as f:
            json.dump(report_data, f, indent=2)

        # Update metadata for success
        meta_entry["status"] = "success"
        meta_entry["fps"] = report_data.get("fps", 0.0)
        meta_entry["duration_sec"] = report_data.get("total_time_sec", 0.0)
        meta_entry["cadence_spm"] = report_data.get("cadence_steps_per_min", 0.0)
        meta_entry["mean_confidence"] = report_data.get("tracking", {}).get("mean", 0.0)

        # Write meta.json
        with open(session_dir / "meta.json", "w", encoding="utf-8") as f:
            json.dump(meta_entry, f, indent=2)

        # Update root index.json
        update_index(meta_entry)

        # One-line summary print
        conf_pct = f"{meta_entry['mean_confidence'] * 100:.1f}%" if meta_entry["mean_confidence"] else "N/A"
        print(f"✅ [{session_id}] Patient: {patient_name} | Duration: {meta_entry['duration_sec']}s | Cadence: {meta_entry['cadence_spm']} spm | Confidence: {conf_pct} | Report: {meta_entry['report_docx']}")

    except Exception as e:
        err_msg = str(e)
        meta_entry["status"] = "failed"
        meta_entry["error"] = err_msg

        # Write meta.json on failure
        with open(session_dir / "meta.json", "w", encoding="utf-8") as f:
            json.dump(meta_entry, f, indent=2)

        update_index(meta_entry)

        print(f"❌ [{session_id}] Processing Failed: {err_msg}")

    return meta_entry


# ── Batch Processing ──────────────────────────────────────────────────────────

def process_batch(folder_path: Path):
    """Process all videos in folder_path skipping ones already in index.json."""
    folder_path = Path(folder_path).resolve()
    if not folder_path.exists() or not folder_path.is_dir():
        print(f"❌ Error: Batch folder path does not exist: {folder_path}")
        return

    valid_exts = {".mp4", ".mov", ".avi", ".mkv"}
    video_files = [
        f for f in folder_path.iterdir()
        if f.is_file() and f.suffix.lower() in valid_exts
    ]

    if not video_files:
        print(f"No video files found in {folder_path}")
        return

    index = load_index()
    processed_filenames = {
        entry.get("video_filename") for entry in index
        if entry.get("status") == "success"
    }

    print(f"📂 Found {len(video_files)} video file(s) in {folder_path.name}")
    results = []

    for vid in video_files:
        if vid.name in processed_filenames:
            print(f"⏭ Skipping already processed video: {vid.name}")
            continue

        res = process_single_video(vid)
        results.append(res)

    print("\n" + "=" * 75)
    print(f"{'FILENAME':<25} | {'SESSION ID':<18} | {'STATUS':<8} | {'CADENCE':<8} | {'CONFIDENCE':<10}")
    print("=" * 75)
    for r in results:
        conf_str = f"{r.get('mean_confidence', 0) * 100:.1f}%" if r.get('mean_confidence') else "N/A"
        cad_str = f"{r.get('cadence_spm', 0)} spm"
        print(f"{r['video_filename']:<25} | {r['session_id']:<18} | {r['status']:<8} | {cad_str:<8} | {conf_str:<10}")
    print("=" * 75)


# ── Cross-Session Trend Summary ───────────────────────────────────────────────

def generate_trend_report(since_date: str = None):
    """Generate cross-session trend report and line plots across sessions."""
    index = load_index()
    if not index:
        print("No sessions recorded in index.json yet.")
        return

    valid_sessions = [s for s in index if s.get("status") == "success"]

    if since_date:
        valid_sessions = [s for s in valid_sessions if s.get("date", "") >= since_date]

    if not valid_sessions:
        print(f"No successful sessions found since {since_date or 'beginning'}.")
        return

    valid_sessions.sort(key=lambda s: (s.get("date", ""), s.get("session_number", 0)))

    print(f"📈 Generating Trend Report for {len(valid_sessions)} session(s)...")

    labels = [f"{s.get('patient_name', 'Session')} ({s['session_id']})" for s in valid_sessions]
    cadences = [s.get("cadence_spm", 0) for s in valid_sessions]
    confidences = [(s.get("mean_confidence", 0) or 0) * 100 for s in valid_sessions]
    durations = [s.get("duration_sec", 0) for s in valid_sessions]

    fig, axes = plt.subplots(3, 1, figsize=(10, 8), sharex=True)
    fig.suptitle("Underwater Gait Analysis - Recovery Trend Report", fontsize=14)

    axes[0].plot(labels, cadences, "o-", color="#0071e3", linewidth=2)
    axes[0].set_title("Cadence Over Sessions (steps/min)")
    axes[0].set_ylabel("Steps / Min")
    axes[0].grid(True, alpha=0.3)

    axes[1].plot(labels, confidences, "o-", color="#30d158", linewidth=2)
    axes[1].set_title("Tracking Confidence (%)")
    axes[1].set_ylabel("Confidence (%)")
    axes[1].set_ylim([0, 100])
    axes[1].grid(True, alpha=0.3)

    axes[2].plot(labels, durations, "o-", color="#ff9f0a", linewidth=2)
    axes[2].set_title("Session Duration (seconds)")
    axes[2].set_ylabel("Duration (s)")
    axes[2].tick_params(axis="x", rotation=45)
    axes[2].grid(True, alpha=0.3)

    plt.tight_layout()
    trend_png = SESSIONS_DIR / "trend_chart.png"
    plt.savefig(trend_png, dpi=150, bbox_inches="tight")
    plt.close()

    doc = Document()
    title = doc.add_heading("Cross-Session Recovery Trend Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Sessions Included: {len(valid_sessions)}")
    if since_date:
        doc.add_paragraph(f"Since Date: {since_date}")

    doc.add_heading("1. Progress Summary Table", level=1)
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Patient Name"
    hdr[1].text = "Session ID"
    hdr[2].text = "Duration (s)"
    hdr[3].text = "Cadence (spm)"
    hdr[4].text = "Confidence (%)"

    for s in valid_sessions:
        r = table.add_row().cells
        r[0].text = s.get("patient_name", "Unknown Patient")
        r[1].text = s["session_id"]
        r[2].text = str(s.get("duration_sec", 0))
        r[3].text = str(s.get("cadence_spm", 0))
        conf_val = f"{s.get('mean_confidence', 0) * 100:.1f}%" if s.get('mean_confidence') else "N/A"
        r[4].text = conf_val

    doc.add_heading("2. Recovery Trend Visualization", level=1)
    if trend_png.exists():
        doc.add_picture(str(trend_png), width=Inches(6.0))

    trend_docx = SESSIONS_DIR / "trend_report.docx"
    doc.save(str(trend_docx))

    print(f"✅ Trend report generated: {trend_docx}")


# ── Main CLI Parser ───────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Automated Session-Based Gait Report Generation")
    parser.add_argument("--video", type=str, help="Path to input raw clip (.mp4, .mov)")
    parser.add_argument("--patient", type=str, default="Unknown Patient", help="Patient Name")
    parser.add_argument("--date", type=str, help="Recorded Date (YYYY-MM-DD)")
    parser.add_argument("--time", type=str, help="Recorded Time (HH:MM)")
    parser.add_argument("--batch", type=str, help="Path to directory containing video clips for batch processing")
    parser.add_argument("--trend", action="store_true", help="Generate cross-session trend summary report")
    parser.add_argument("--since", type=str, help="Start date filter for --trend (YYYY-MM-DD)")

    args = parser.parse_args()

    if args.video:
        process_single_video(
            Path(args.video), 
            date_str=args.date, 
            patient_name=args.patient,
            time_str=args.time
        )
    elif args.batch:
        process_batch(Path(args.batch))
    elif args.trend:
        generate_trend_report(since_date=args.since)
    else:
        parser.print_help()


if __name__ == "__main__":
    main()
